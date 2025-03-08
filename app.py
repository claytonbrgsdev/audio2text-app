import os
import streamlit as st
import whisper
import subprocess
import tempfile
import numpy as np
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document
from docx.shared import Pt, Inches

try:
    from pyannote.audio import Pipeline
    import torch
except ImportError:
    pass  # We'll handle if not installed

# --- Language dictionaries for UI text ---
LANG_DICT = {
    "en": {
        "title": "Audio Transcription & Diarization App",
        "upload_desc": "Upload an audio file (up to 1GB) to transcribe using OpenAI Whisper. Optionally enable speaker diarization via pyannote.audio 3.1.",
        "upload_label": "Upload Audio File",
        "select_model_label": "Select Whisper Model",
        "select_model_help": (
            "Tiny/Base/Small are faster but less accurate. Medium/Large-V2 are slower but more accurate. "
            "Choose based on your hardware and desired accuracy."
        ),
        "diarization_label": "Enable Diarization",
        "diarization_help": "Requires huggingface-cli login + gating for pyannote/speaker-diarization-3.1",
        "transcribe_button": "Transcribe Audio",
        "warn_upload": "Please upload an audio file first.",
        "transcription_done": "Transcription complete!",
        "diarization_done": "Diarization complete!",
        "diarization_results": "Diarization Results",
        "speaker_labeled_transcript": "Speaker-Labeled Transcript",
        "transcript_no_diar": "Transcription (No Diarization)",
        "rename_speakers_header": "Rename Speakers",
        "rename_speakers_info": "You can replace default labels (e.g. SPEAKER_00) with custom names (e.g. Gabriela).",
        "apply_rename_button": "Apply Speaker Renaming",
        "renamed_transcript": "Transcript with Renamed Speakers",
        "download_pdf": "Download as PDF",
        "download_docx": "Download as DOCX"
    },
    "pt": {
        "title": "Aplicativo de Transcrição e Diarização de Áudio",
        "upload_desc": "Envie um arquivo de áudio (até 1GB) para transcrever usando o OpenAI Whisper. Opcionalmente, ative a diarização via pyannote.audio 3.1.",
        "upload_label": "Envie seu arquivo de áudio",
        "select_model_label": "Selecione o Modelo Whisper",
        "select_model_help": (
            "Tiny/Base/Small são mais rápidos porém menos precisos. Medium/Large-V2 são mais lentos mas mais precisos. "
            "Escolha de acordo com seu hardware e a acurácia desejada."
        ),
        "diarization_label": "Ativar Diarização",
        "diarization_help": "Requer login no huggingface-cli + acesso ao pyannote/speaker-diarization-3.1",
        "transcribe_button": "Transcrever Áudio",
        "warn_upload": "Por favor, envie um arquivo de áudio primeiro.",
        "transcription_done": "Transcrição concluída!",
        "diarization_done": "Diarização concluída!",
        "diarization_results": "Resultados de Diarização",
        "speaker_labeled_transcript": "Transcrição com Falantes Identificados",
        "transcript_no_diar": "Transcrição (Sem Diarização)",
        "rename_speakers_header": "Renomear Falantes",
        "rename_speakers_info": "Você pode substituir rótulos padrão (ex. SPEAKER_00) por nomes personalizados (ex. Gabriela).",
        "apply_rename_button": "Aplicar Renomeação",
        "renamed_transcript": "Transcrição com Falantes Renomeados",
        "download_pdf": "Baixar como PDF",
        "download_docx": "Baixar como DOCX"
    }
}

# --- Utility Functions ---

def get_audio_duration(file_path):
    result = subprocess.run(
        ["ffprobe", "-v", "error", "-show_entries", "format=duration",
         "-of", "default=noprint_wrappers=1:nokey=1", file_path],
        stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True
    )
    return float(result.stdout.strip())

def segment_audio(file_path, max_duration=600, max_file_size=100*1024*1024):
    total_duration = get_audio_duration(file_path)
    file_size = os.path.getsize(file_path)
    if total_duration <= max_duration and file_size <= max_file_size:
        return [(file_path, 0.0)]
    
    num_segments = int(np.ceil(total_duration / max_duration))
    segment_duration = total_duration / num_segments
    
    segments = []
    for i in range(num_segments):
        chunk_start = i * segment_duration
        current_duration = min(segment_duration, total_duration - chunk_start)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as temp:
            segment_path = temp.name
        ffmpeg_cmd = [
            "ffmpeg", "-y", "-i", file_path,
            "-ss", str(chunk_start), "-t", str(current_duration),
            "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1", segment_path
        ]
        subprocess.run(ffmpeg_cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        segments.append((segment_path, chunk_start))
    return segments

def transcribe_audio(file_path, model, progress_bar):
    segments = segment_audio(file_path)
    # If only one segment
    if len(segments) == 1:
        with st.spinner("Transcribing audio..."):
            result = model.transcribe(file_path, fp16=False)
            progress_bar.progress(100)
        return result["text"].strip(), result.get("segments", [])
    else:
        full_text = ""
        transcription_segments = []
        total_segments = len(segments)
        for i, (seg, chunk_start) in enumerate(segments):
            with st.spinner(f"Transcribing segment {i+1}/{total_segments}..."):
                result = model.transcribe(seg, fp16=False)
                seg_text = result["text"].strip()
                full_text += seg_text + " "
                seg_segments = result.get("segments", [])
                for s in seg_segments:
                    s["start"] += chunk_start
                    s["end"] += chunk_start
                    transcription_segments.append(s)
                progress = int(((i + 1) / total_segments) * 100)
                progress_bar.progress(progress)
            if seg != file_path:
                os.remove(seg)
        return full_text.strip(), transcription_segments

def convert_to_wav(input_path, sr=16000):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp:
        wav_path = tmp.name
    ffmpeg_cmd = [
        "ffmpeg", "-y", "-i", input_path,
        "-ar", str(sr), "-ac", "1", wav_path
    ]
    subprocess.run(ffmpeg_cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    return wav_path

def assign_speakers_to_transcription(transcription_segments, diarization_results):
    if not transcription_segments:
        return ""
    for seg in transcription_segments:
        midpoint = (seg["start"] + seg["end"]) / 2
        assigned_speaker = None
        for d in diarization_results:
            d_start = float(d["Start (sec)"])
            d_end = float(d["End (sec)"])
            if d_start <= midpoint <= d_end:
                assigned_speaker = d["Speaker"]
                break
        seg["speaker"] = assigned_speaker if assigned_speaker else "Unknown"
    
    merged_lines = []
    current_speaker = None
    current_text = ""
    for seg in transcription_segments:
        speaker = seg.get("speaker", "Unknown")
        text = seg.get("text", "").strip()
        if current_speaker != speaker:
            if current_speaker is not None:
                merged_lines.append(f"{current_speaker}: {current_text.strip()}")
            current_speaker = speaker
            current_text = text + " "
        else:
            current_text += text + " "
    if current_text:
        merged_lines.append(f"{current_speaker}: {current_text.strip()}")
    return "\n\n".join(merged_lines)

def generate_pdf(text):
    import tempfile
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    temp_pdf.close()
    
    pdf = SimpleDocTemplate(temp_pdf.name, pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    style = ParagraphStyle("Normal", fontName="Helvetica", fontSize=11, leading=14)
    
    elements = [
        Paragraph("Audio Transcription", styles['Title']),
        Paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']),
        Spacer(1, 20)
    ]
    for para in text.split("\n\n"):
        elements.append(Paragraph(para, style))
        elements.append(Spacer(1, 12))
    
    pdf.build(elements)
    with open(temp_pdf.name, "rb") as f:
        pdf_bytes = f.read()
    os.remove(temp_pdf.name)
    return pdf_bytes

def generate_docx(text):
    import tempfile
    doc = Document()
    doc.add_heading("Audio Transcription", 0)
    doc.add_paragraph(f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    for para in text.split("\n\n"):
        doc.add_paragraph(para)
    temp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    temp_docx.close()
    doc.save(temp_docx.name)
    with open(temp_docx.name, "rb") as f:
        docx_bytes = f.read()
    os.remove(temp_docx.name)
    return docx_bytes

# --- Main App ---
def main():
    # Language Switch
    lang_choice = st.radio("Language / Idioma", ["English", "Português"], index=0)
    lang_key = "en" if lang_choice == "English" else "pt"
    text = LANG_DICT[lang_key]
    
    st.title(text["title"])
    st.write(text["upload_desc"])
    
    # Initialize session state
    if "transcription_done" not in st.session_state:
        st.session_state.transcription_done = False
        st.session_state.transcription_text = ""
        st.session_state.transcription_segments = []
        st.session_state.diarization_results = []
        st.session_state.final_transcript = ""
        st.session_state.speaker_map = {}
    
    uploaded_file = st.file_uploader(text["upload_label"], type=["mp3", "wav", "flac", "m4a", "ogg"])
    if uploaded_file is not None:
        st.audio(uploaded_file)
    
    model_choice = st.selectbox(
        text["select_model_label"],
        ["tiny", "base", "small", "medium", "large-v2", "large-v3"],
        help=text["select_model_help"]
    )
    
    enable_diarization = st.checkbox(text["diarization_label"], value=False, help=text["diarization_help"])
    
    if st.button(text["transcribe_button"]):
        if not uploaded_file:
            st.warning(text["warn_upload"])
        else:
            # Clear old data
            st.session_state.transcription_done = False
            st.session_state.transcription_text = ""
            st.session_state.transcription_segments = []
            st.session_state.diarization_results = []
            st.session_state.final_transcript = ""
            st.session_state.speaker_map = {}
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp:
                tmp.write(uploaded_file.getvalue())
                file_path = tmp.name
            
            progress_bar = st.progress(0)
            
            # Load model with spinner
            with st.spinner("Loading Whisper model..."):
                model = whisper.load_model(model_choice)
            
            # Transcribe
            transcription_text, transcription_segments = transcribe_audio(file_path, model, progress_bar)
            st.session_state.transcription_text = transcription_text
            st.session_state.transcription_segments = transcription_segments
            st.session_state.final_transcript = transcription_text
            st.session_state.transcription_done = True
            
            # Diarization if checked
            if enable_diarization and "Pipeline" in globals():
                with st.spinner("Performing diarization. This may take some time..."):
                    try:
                        device = "cpu"
                        if "torch" in globals() and torch.backends.mps.is_available():
                            device = "mps"
                        pipeline = Pipeline.from_pretrained(
                            "pyannote/speaker-diarization-3.1",
                            use_auth_token=True
                        )
                        try:
                            pipeline.to(device)
                        except:
                            pass
                        diarization_wav = convert_to_wav(file_path, sr=16000)
                        diarization = pipeline(diarization_wav)
                        diarization_results = []
                        for segment, _, speaker in diarization.itertracks(yield_label=True):
                            diarization_results.append({
                                "Start (sec)": f"{segment.start:.2f}",
                                "End (sec)": f"{segment.end:.2f}",
                                "Speaker": speaker
                            })
                        st.session_state.diarization_results = diarization_results
                        os.remove(diarization_wav)
                    except Exception as e:
                        st.error(f"Error during diarization: {e}")
            
            os.remove(file_path)
            st.success(text["transcription_done"])
    
    # Display results
    if st.session_state.transcription_done:
        if st.session_state.diarization_results:
            st.subheader(text["diarization_results"])
            st.table(st.session_state.diarization_results)
            
            if st.session_state.transcription_segments:
                merged_transcript = assign_speakers_to_transcription(
                    st.session_state.transcription_segments,
                    st.session_state.diarization_results
                )
                st.session_state.final_transcript = merged_transcript
            else:
                st.session_state.final_transcript = st.session_state.transcription_text
            
            st.subheader(text["speaker_labeled_transcript"])
            st.text_area("Transcript", value=st.session_state.final_transcript, height=300)
        else:
            st.subheader(text["transcript_no_diar"])
            st.text_area("Transcript", value=st.session_state.transcription_text, height=300)
            st.session_state.final_transcript = st.session_state.transcription_text
        
        # Speaker rename
        speaker_labels = sorted({d["Speaker"] for d in st.session_state.diarization_results})
        if speaker_labels:
            st.markdown("---")
            st.markdown(f"### {text['rename_speakers_header']}")
            st.info(text["rename_speakers_info"])
            
            if not st.session_state.speaker_map:
                for lbl in speaker_labels:
                    st.session_state.speaker_map[lbl] = lbl
            
            for lbl in speaker_labels:
                st.session_state.speaker_map[lbl] = st.text_input(
                    f"{lbl}:", value=st.session_state.speaker_map[lbl], key=f"rename_{lbl}"
                )
            
            if st.button(text["apply_rename_button"]):
                updated_text = st.session_state.final_transcript
                for lbl in speaker_labels:
                    new_name = st.session_state.speaker_map[lbl]
                    updated_text = updated_text.replace(lbl + ":", new_name + ":")
                st.session_state.final_transcript = updated_text
                st.success("Done!")
            
            st.subheader(text["renamed_transcript"])
            st.text_area("Transcript (Renamed)", value=st.session_state.final_transcript, height=300)
        
        # Download
        pdf_bytes = generate_pdf(st.session_state.final_transcript)
        docx_bytes = generate_docx(st.session_state.final_transcript)
        st.download_button(
            text["download_pdf"], data=pdf_bytes,
            file_name=f"transcription_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
            mime="application/pdf"
        )
        st.download_button(
            text["download_docx"], data=docx_bytes,
            file_name=f"transcription_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()